"""
DOCX Processor - Extract text and tables from Word documents
Supports .docx format (Office Open XML)
"""
#backend/app/core/processors/docx_processor.py


import time
from typing import List, Dict, Any
import pandas as pd
import logging

from .base import BaseProcessor, ProcessingResult
from app.config import settings

logger = logging.getLogger(__name__)


class DOCXProcessor(BaseProcessor):
    """Process Word documents and extract text and tables"""
    
    def __init__(self):
        super().__init__()
    
    def process(self, file_path: str, **kwargs) -> ProcessingResult:
        """
        Process DOCX file and extract content
        
        Args:
            file_path: Path to DOCX file
            **kwargs:
                - extract_text: bool (default: True)
                - extract_tables: bool (default: True)
                - include_formatting: bool (default: False)
        
        Returns:
            ProcessingResult with extracted data
        """
        start_time = time.time()
        
        try:
            # Validate file
            self.validate_file(file_path)
            
            # Get file metadata
            metadata = self.get_file_metadata(file_path)
            
            # Extract options
            extract_text = kwargs.get('extract_text', True)
            extract_tables = kwargs.get('extract_tables', True)
            include_formatting = kwargs.get('include_formatting', False)
            
            # Import python-docx
            from docx import Document
            
            # Load document
            doc = Document(file_path)
            
            # Extract content
            text_content = ""
            dataframes = []
            warnings = []
            
            if extract_text:
                text_content = self._extract_text(doc, include_formatting)
            
            if extract_tables:
                tables = self._extract_tables(doc)
                
                for i, table_df in enumerate(tables):
                    if not table_df.empty:
                        cleaned_df = self.clean_dataframe(table_df)
                        dataframes.append(cleaned_df)
            
            # Extract document metadata
            doc_metadata = self._extract_document_metadata(doc)
            metadata.update(doc_metadata)
            
            # Add document statistics
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            # Create result
            result = ProcessingResult(
                success=True,
                file_path=file_path,
                file_type='docx',
                dataframes=dataframes,
                text_content=text_content,
                metadata=metadata,
                processing_time=time.time() - start_time,
                warnings=warnings,
                sheet_names=[f"Table_{i+1}" for i in range(len(dataframes))]
            )
            
            self.log_processing_info(result)
            return result
        
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            return ProcessingResult(
                success=False,
                file_path=file_path,
                file_type='docx',
                error_message=str(e),
                processing_time=time.time() - start_time
            )
    
    def _extract_text(self, doc, include_formatting: bool = False) -> str:
        """
        Extract text from document
        
        Args:
            doc: python-docx Document object
            include_formatting: Whether to include formatting info
        
        Returns:
            Extracted text
        """
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if include_formatting:
                    # Include style information
                    style = paragraph.style.name if paragraph.style else "Normal"
                    text_parts.append(f"[{style}] {paragraph.text}")
                else:
                    text_parts.append(paragraph.text)
        
        return "\n".join(text_parts)
    
    def _extract_tables(self, doc) -> List[pd.DataFrame]:
        """
        Extract tables from document
        
        Args:
            doc: python-docx Document object
        
        Returns:
            List of DataFrames containing tables
        """
        dataframes = []
        
        for table in doc.tables:
            # Extract table data
            data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                data.append(row_data)
            
            if data and len(data) > 1:
                # Use first row as header if it looks like a header
                if self._is_likely_header(data[0], data[1:]):
                    df = pd.DataFrame(data[1:], columns=data[0])
                else:
                    df = pd.DataFrame(data)
                
                dataframes.append(df)
        
        return dataframes
    
    def _is_likely_header(
        self,
        first_row: List[str],
        remaining_rows: List[List[str]]
    ) -> bool:
        """
        Determine if first row is likely a header
        
        Args:
            first_row: First row data
            remaining_rows: Rest of the table data
        
        Returns:
            True if first row is likely a header
        """
        # Check if first row has unique values
        if len(set(first_row)) != len(first_row):
            return False
        
        # Check if first row is all strings and other rows have numbers
        if not remaining_rows:
            return True
        
        first_row_is_text = all(
            not self._is_number(cell) for cell in first_row if cell
        )
        
        # Heuristic: headers often have shorter strings and fewer blanks than body rows
        avg_head_len = sum(len(str(c)) for c in first_row) / len(first_row) if first_row else 0
        
        body_stats = []
        for row in remaining_rows[:5]: # Check first 5 rows
            avg_len = sum(len(str(c)) for c in row) / len(row) if row else 0
            blanks = sum(1 for c in row if not str(c).strip())
            body_stats.append((avg_len, blanks))
            
        other_rows_have_numbers = any(
            any(self._is_number(cell) for cell in row)
            for row in remaining_rows
        )
        
        # If headers are shorter on average and body has numbers, it's likely a header
        is_shorter = any(avg_head_len < stats[0] * 0.8 for stats in body_stats) if body_stats else False
        
        return first_row_is_text and (other_rows_have_numbers or is_shorter)
    
    def _is_number(self, s: str) -> bool:
        """Check if string represents a number"""
        try:
            float(s.replace(',', ''))
            return True
        except (ValueError, AttributeError):
            return False
    
    def _extract_document_metadata(self, doc) -> Dict[str, Any]:
        """
        Extract document properties/metadata
        
        Args:
            doc: python-docx Document object
        
        Returns:
            Dictionary with metadata
        """
        metadata = {}
        
        try:
            core_props = doc.core_properties
            
            metadata['author'] = core_props.author
            metadata['title'] = core_props.title
            metadata['subject'] = core_props.subject
            metadata['keywords'] = core_props.keywords
            metadata['comments'] = core_props.comments
            metadata['created'] = str(core_props.created) if core_props.created else None
            metadata['modified'] = str(core_props.modified) if core_props.modified else None
            metadata['last_modified_by'] = core_props.last_modified_by
            metadata['revision'] = core_props.revision
        
        except Exception as e:
            self.logger.warning(f"Could not extract document metadata: {str(e)}")
        
        return metadata
    
    def extract_headings(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract document structure (headings)
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            List of headings with their levels and text
        """
        try:
            from docx import Document
            doc = Document(file_path)
            
            headings = []
            
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.split()[-1])
                    headings.append({
                        'level': level,
                        'text': paragraph.text,
                        'style': paragraph.style.name
                    })
            
            return headings
        
        except Exception as e:
            self.logger.error(f"Error extracting headings: {str(e)}")
            return []
    
    def extract_images_info(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract information about images in document
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            List of image information
        """
        try:
            from docx import Document
            doc = Document(file_path)
            
            images_info = []
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images_info.append({
                        'target': rel.target_ref,
                        'type': rel.reltype
                    })
            
            return images_info
        
        except Exception as e:
            self.logger.error(f"Error extracting image info: {str(e)}")
            return []
    
    def get_document_statistics(self, file_path: str) -> Dict[str, Any]:
        """
        Get comprehensive document statistics
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Dictionary with statistics
        """
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Count words
            total_words = 0
            total_chars = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                total_words += len(text.split())
                total_chars += len(text)
            
            return {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'word_count': total_words,
                'character_count': total_chars,
                'page_estimate': total_words // 250  # Rough estimate
            }
        
        except Exception as e:
            self.logger.error(f"Error getting statistics: {str(e)}")
            return {}